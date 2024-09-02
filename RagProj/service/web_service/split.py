import os

from pypdf import PdfReader
from docx import Document
import re
from tqdm import tqdm


def split_content_to_parse(content, max_length):
    sentences = re.split(r"([。！？；.!?;])", content)
    sentences.append("")
    sentences = ["".join(i) for i in zip(sentences[0::2], sentences[1::2])]
    if sentences[-1] == "":
        sentences.pop(-1)
    all_paras = []
    all_sentences_num_in_paras = []
    paras = []
    sentences_num_in_paras = 0
    sentences_num = len(sentences)
    for idx, sen in enumerate(sentences):
        if len("".join(paras)) <= max_length:
            paras.append(sen)
            sentences_num_in_paras += 1
        if len("".join(paras)) > max_length:
            if sentences_num_in_paras > 1:
                all_paras.append("".join(paras[:-1]))
                all_sentences_num_in_paras.append(sentences_num_in_paras - 1)
                paras = []
                sentences_num_in_paras = 1
                paras.append(sen)
            else:
                all_paras.append("".join(paras))
                all_sentences_num_in_paras.append(sentences_num_in_paras)
                paras = []
                sentences_num_in_paras = 0
        if idx == sentences_num - 1 and sentences_num_in_paras >= 1:
            all_paras.append("".join(paras))
            all_sentences_num_in_paras.append(sentences_num_in_paras)
    return all_paras, all_sentences_num_in_paras


def load_file(file_path, max_para_length=512):
    max_para_length = max_para_length

    def _get_pdf_lines(pdf_path):
        reader = PdfReader(pdf_path)
        number_of_pages = len(reader.pages)
        all_paras = []
        print("Start loading pdf")
        result = []
        for i in tqdm(range(number_of_pages)):
            page = reader.pages[i]
            all_lines = page.extract_text()
            paras, _ = split_content_to_parse(all_lines, max_para_length)
            all_paras += paras
        return all_paras

    def _get_doc_lines(doc_path):
        doc = Document(doc_path)
        all_paras = []
        print("Start loading doc")
        for paragraph in tqdm(doc.paragraphs):
            paras, _ = split_content_to_parse(paragraph.text, max_para_length)
            all_paras += paras
        return all_paras

    def load_files(file_path):
        para = []
        if file_path.endswith("pdf"):
            para = _get_pdf_lines(file_path)
        elif file_path.endswith("docx"):
            para = _get_doc_lines(file_path)
        return para

    return load_files(file_path=file_path)


class LoadFileService:
    def __init__(self, max_para_length=512):
        self.max_para_length = max_para_length

    def load_files(self, file_path):
        para = []
        if file_path.endswith("pdf"):
            para = self._get_pdf_lines(file_path)
        elif file_path.endswith("docx"):
            para = self._get_doc_lines(file_path)
        return para

    def _get_pdf_lines(self, pdf_path):
        reader = PdfReader(pdf_path)
        number_of_pages = len(reader.pages)
        all_paras = []
        print("Start loading pdf")
        result = []
        for i in tqdm(range(number_of_pages)):
            page = reader.pages[i]
            all_lines = page.extract_text()
            paras, _ = split_content_to_parse(all_lines, self.max_para_length)
            all_paras += paras
        return all_paras

    def _get_doc_lines(self, doc_path):
        doc = Document(doc_path)
        all_paras = []
        print("Start loading doc")
        for paragraph in tqdm(doc.paragraphs):
            paras, _ = split_content_to_parse(paragraph.text, self.max_para_length)
            all_paras += paras
        return all_paras
